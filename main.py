# -*- coding:utf-8 -*-

import sys
from PyQt5 import QtWidgets
from common.publish_helper import build_output, get_similarity_threshold, set_similarity_threshold
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from parser import BibParser, MultiLinesParser
from setting_dialog import SettingDialog
reload(sys)
sys.setdefaultencoding('utf-8')


class AppWindow(QtWidgets.QMainWindow):

    def __init__(self, parent=None):
        super(AppWindow, self).__init__(parent)
        self.create_main_ui()
        self.bind_actions()
        self.resize(1024, 768)

    # UI definition
    def create_main_ui(self):
        self.main_frame = QtWidgets.QWidget()
        main_v_layout = QtWidgets.QVBoxLayout()
        self.origin_text = QtWidgets.QTextEdit()
        main_v_layout.addWidget(self.origin_text)
        button_h_layout = QtWidgets.QHBoxLayout()
        self.btn_select_file = QtWidgets.QPushButton('导入文件')
        self.btn_config_label = QtWidgets.QPushButton('配置转换规则')
        self.btn_set_threshold = QtWidgets.QPushButton('设置相似度阈值')
        self.btn_begin_convert = QtWidgets.QPushButton('开始转换')
        self.btn_export_file = QtWidgets.QPushButton('结果导出文件')
        button_h_layout.addWidget(self.btn_select_file)
        button_h_layout.addWidget(self.btn_config_label)
        button_h_layout.addWidget(self.btn_set_threshold)
        button_h_layout.addWidget(self.btn_begin_convert)
        button_h_layout.addWidget(self.btn_export_file)
        main_v_layout.addLayout(button_h_layout)
        self.result_text = QtWidgets.QTextEdit()
        main_v_layout.addWidget(self.result_text)
        self.main_frame.setLayout(main_v_layout)
        self.setCentralWidget(self.main_frame)

    def bind_actions(self):
        self.btn_select_file.clicked.connect(self.on_btn_select_file_click)
        self.btn_begin_convert.clicked.connect(self.on_btn_begin_convert_click)
        self.btn_config_label.clicked.connect(self.on_btn_config_label_click)
        self.btn_set_threshold.clicked.connect(self.on_btn_set_threshold_click)
        self.btn_export_file.clicked.connect(self.on_btn_export_file_click)

    # event handler
    def on_btn_select_file_click(self):
        filenames, _ = QtWidgets.QFileDialog.getOpenFileNames(self, "选择文件")
        if not filenames:
            return
        content = self.origin_text.toPlainText()
        for filename in filenames:
            try:
                with open(filename, 'r') as f:
                    content += f.read()
            except:
                pass
        self.origin_text.setText(content)

    def on_btn_begin_convert_click(self):
        content = self.origin_text.toPlainText()
        # parser = BibParser()
        # result = parser.parse_string(content)
        # if result is None:
        parser = MultiLinesParser()
        result = parser.parse_string(content)
        output_content = build_output(result)
        self.result_text.setText(output_content)
        self.add_message_box('转换完成')

    def on_btn_config_label_click(self):
        dialog = SettingDialog()
        if dialog.exec_():
            dialog.config_mata_data_base()

    def on_btn_set_threshold_click(self):
        threshold, ok = QtWidgets.QInputDialog.getDouble(self, '设置相似度阈值', '设置值',
                                                         get_similarity_threshold(), 0, 1.0, 2)
        if ok:
            set_similarity_threshold(threshold)

    def add_message_box(self, message, title='提示'):
        msg = QtWidgets.QMessageBox()
        msg.setText(message)
        msg.setWindowTitle(title)
        msg.exec_()

    def on_btn_export_file_click(self):
        message = self.result_text.toPlainText()
        save_filename, ext = QtWidgets.QFileDialog.getSaveFileName(self, 'Save File', '', '*.docx')
        if not save_filename:
            return
        document = Document()
        p = document.add_paragraph('')
        for c in message:
            run = p.add_run(c)
            run.font.size = Pt(12)
            if c.isalpha() or c.isdigit():
                run.font.name = 'Time New Roman'
            else:
                run.font.name = u'宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.save(save_filename)


def main():
    app = QtWidgets.QApplication(sys.argv)
    window = AppWindow()
    window.show()
    app.exec_()


if __name__ == '__main__':
    main()
